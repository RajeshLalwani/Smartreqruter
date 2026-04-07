import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG = lambda name: os.path.join(BASE_DIR, 'Diagrams', name) if os.path.exists(os.path.join(BASE_DIR, 'Diagrams', name)) else os.path.join(BASE_DIR, name)

doc = Document()

# ── PAGE MARGINS ─────────────────────────────
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ── STYLE HELPERS ─────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def add_image(path, width=5.5, caption=None):
    full = IMG(path)
    if os.path.exists(full):
        doc.add_picture(full, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(10)
    else:
        p = doc.add_paragraph(f"[Image placeholder: {caption or path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_data_dictionary_table(title, rows):
    # Title
    p = doc.add_paragraph()
    r = p.add_run(f"Table Name: {title}")
    r.bold = True
    r.underline = True
    
    headers = ["Sr. No", "Field Name", "Data Type", "Constraints"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    
    # Format header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr[i], '3B78B0') # Header color
    
    # Format rows with alternating colors
    colors = ['DDEBF7', 'BDD7EE']
    for row_idx, row_data in enumerate(rows):
        row = t.add_row().cells
        bg_color = colors[row_idx % 2]
        for idx, val in enumerate(row_data):
            row[idx].text = str(val)
            set_cell_background(row[idx], bg_color)
    
    # Column width setting
    widths = [0.8, 2.0, 1.5, 2.0]
    for row in t.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    doc.add_paragraph()

# ═══════════════════════════════════════════
#  TITLE PAGE (Modified to align with tech elecon constraints if needed)
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph("SMARTRECRUIT AI RECRUITMENT PLATFORM")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].bold = True
t.runs[0].font.size = Pt(22)
t.runs[0].font.color.rgb = RGBColor(0, 70, 127)

doc.add_paragraph()
sub = doc.add_paragraph("A Next-Generation AI-Powered Talent Acquisition Ecosystem")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
sub.runs[0].font.size = Pt(14)

doc.add_paragraph()
para("Project Report submitted as partial fulfilment of the degree of\nMaster of Computer Applications (MCA)", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("Submitted by\nCandidate Name\nSeat No.: XXXX", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("Post Graduate Department of Computer Science and Technology\nSardar Patel University\nVallabh Vidyanagar", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("Submitted to", align=WD_ALIGN_PARAGRAPH.CENTER)
para("SARDAR PATEL UNIVERSITY\nVallabh Vidyanagar\nApril 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CERTIFICATE
# ═══════════════════════════════════════════
heading("CERTIFICATE", 1)
para(
    "This is to certify that Candidate Name of Master of Computer Applications (MCA) "
    "IV semester has worked on the project entitled SmartRecruit AI Recruitment Platform satisfactory "
    "towards the partial fulfilment of the degree of Master of Computer Applications (MCA) during the "
    "final semester at the Post Graduate Department of Computer Science and Technology, Sardar Patel "
    "University, Vallabh Vidyanagar, Gujarat, India."
)
doc.add_paragraph()
para("Date of Submission 25th April, 2026")
doc.add_paragraph()
para("Internal Project Guide          Head of the Department", bold=True)
doc.add_page_break()

# ═══════════════════════════════════════════
#  ACKNOWLEDGMENT
# ═══════════════════════════════════════════
heading("ACKNOWLEDGMENT", 1)
para(
    "Acknowledging everyone is difficult, but I will still try to express my gratitude to each "
    "and every one who has helped me in the completion of this project.\n\n"
    "I am indebted to Tech Elecon Pvt. Ltd. and my company guides who devoted their precious time to "
    "help me understand the platform that I have been a part of; without their assistance, I would not "
    "have been able to build up a good Application.\n\n"
    "I am grateful to my Internal Project Guide, for their prolonged interest in my work and excellent "
    "guidance. They have been a constant source of motivation for me. By their uncommon promising "
    "demand for quality and their insistence on meeting the deadlines, I was able to accomplish such an "
    "excellent work.\n\n"
    "I express my sincere gratitude to the Head of the Department. They have been very good mentors "
    "who were always available for helping me with any kind of queries that I would have regarding the "
    "topic and anything that would be a hurdle in the development of my application.\n\n"
    "And last but not least, I would like to acknowledge the cooperation extended by all those persons "
    "who have directly or indirectly helped me during the course of this project."
)
doc.add_page_break()


# ═══════════════════════════════════════════
#  INDEX
# ═══════════════════════════════════════════
heading("INDEX", 1)
index_items = [
    "1. INTRODUCTION",
    "2. DEPARTMENT PROFILE",
    "3. COMPANY PROFILE",
    "4. PROJECT OVERVIEW",
    "   4.1 TEAM INFORMATION",
    "   4.2 DIVISION OF WORK",
    "   4.3 PROJECT DEFINITION",
    "   4.4 PURPOSE & OBJECTIVE OF THE PROJECT",
    "   4.5 SCOPE OF THE SYSTEM",
    "   4.6 KEY FEATURES",
    "   4.7 SYSTEM DESIGN TECHNOLOGY STACK",
    "   4.8 LIMITATIONS",
    "5. SYSTEM ANALYSIS",
    "   5.1 EXISTING SYSTEM",
    "   5.2 NEEDS OF SYSTEM",
    "   5.3 PROPOSED SYSTEM",
    "   5.4 MODULES",
    "   5.5 PROJECT MANAGEMENT",
    "   5.6 SYSTEM REQUIREMENT SPECIFICATION",
    "   5.7 SOFTWARE DEVELOPMENT TOOLS",
    "6. SYSTEM DESIGN",
    "   6.1 USE CASE DIAGRAM",
    "   6.2 ACTIVITY DIAGRAMS",
    "   6.3 ER DIAGRAM",
    "   6.4 DFD (DATA FLOW DIAGRAM)",
    "   6.5 SEQUENTIAL DIAGRAM",
    "   6.6 CLASS DIAGRAM",
    "7. DATA DICTIONARY",
    "8. SYSTEM OUTPUT",
    "   8.1 SCREEN SHOT [ADMIN]",
    "   8.2 SCREEN SHOT [USER]",
    "9. TESTING",
    "10. FUTURE ENHANCEMENT",
    "11. CONCLUSION",
    "12. BIBLIOGRAPHY"
]
for item in index_items:
    para(item)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════
heading("1. INTRODUCTION", 1)
para(
    "The SmartRecruit AI Recruitment Platform is a highly advanced, intelligent recruitment portal "
    "designed to fully automate the hiring lifecycle. By taking advantage of Google Gemini Large Language "
    "Models, dynamic skill embeddings, and real-time biometric proctoring, it minimizes the manual overhead "
    "experienced by traditional HR departments.\n\n"
    "In today's fast-paced corporate environment, screening resumes, scheduling technical coding rounds, "
    "and conducting tedious behavioral interviews results in a significant financial and time commitment. "
    "SmartRecruit solves these core inefficiencies through a robust 4-Round AI evaluation pipeline: Resume "
    "Parsing & Match score, Technical Aptitude MCQ generation, Real-time AI Video Interview & Code Exec, and "
    "an AI HR Simulation 'Botanist' round."
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 2: DEPARTMENT PROFILE
# ═══════════════════════════════════════════
heading("2. DEPARTMENT PROFILE", 1)
para("GH Patel PG Department of Computer Science & Technology,")
para("Sardar Patel University, Vallabh Vidyanagar")
para(
    "The Department of Computer Science was established in year 1986 as an institution offering "
    "postgraduate in Computer Science. The computing facility, however, was introduced at the "
    "Sardar Patel University somewhere in 1973-74 at Computer centre of the University with an IBM "
    "1620, a second-generation mainframe computer. The University started one-year Post Graduate "
    "Diploma in Computer Applications (PGDCA) course in 1986, M.Sc. (Computer Science) in 2008, "
    "Ph.D. (Computer Science) in 1990 and Ph.D. (Bioinformatics) in 2011.\n\n"
    "Besides these programs, the University also started computer related short-term programs from "
    "time to time. The computing resources were also used by different postgraduate departments of "
    "the University for their Research Work. The institute took lead in Establishment started a CSI "
    "(Computer Society of India is a national professional body having links with various professional "
    "bodies in India and abroad) in 1985-86. Various workshops, seminars and symposia have been "
    "organized under the CSI banner so far, which have strengthened both the teaching and research "
    "activities of the Department. In the year 1988, Late Shree Goradhanbhai Hathibhai Patel donated "
    "money to the University for a Separate building of the Post Graduate Department of Computer "
    "Science. The Department also recognized by UGC to provide technical training to personnel at various "
    "levels of academia and administration. Several PC based systems had been Installed and major "
    "research work carried out in the areas of CAI (Computer Aided Instruction), KBS (Knowledge Based "
    "Systems) and DSS (Decision Support System). In 1992. The M.Sc. In Computer Science Program was "
    "terminated and intake of the MCA Program was Increased from 30 to 40. Three faculty members were "
    "selected on the basis of merit for the CICC-Japan training in the years 1992, 1994 and 1995.\n\n"
    "Early 2000 was the time when vast technological changes took place in the IT industry all over "
    "the world. The department could meet the technological requirements through support from the UGC, "
    "AICTE and various bodies of the state. A networked laboratory was established at the department, "
    "providing state-of-the-art facilities. The intake to MCA was increased to 100 by introducing "
    "self-financed seats in the year 2000, and a separate self-financed batch of MCA was started. "
    "The department was recognized and offered 'Refresher Courses in Computer Science' for the Computer "
    "Science faculty members in India by the University Grants Commission. In the year 2002, a special "
    "MCA program was introduced for providing lateral entry Into the Second year to computer science "
    "graduates. In a short span, the department grew to the extent that a new building was required. "
    "With one more generous donation from Late Shree Goradhanbhai Hathibhai Patel, grants from the UGC "
    "and he state government and the department with modern infrastructure was constructed. The building "
    "was inaugurated on July 22nd, was then changed to G H Patel Post Graduate Department of Computer "
    "Science and Technology in the year 2002. Subsequently, new powerful servers and several workstations "
    "were added to the laboratory and electives for bioinformatics and wireless networks were added to "
    "the syllabus. The M.Sc. (Bioinformatics) program under the UGC Innovative Program scheme was "
    "introduced in the year 2005."
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 3: COMPANY PROFILE
# ═══════════════════════════════════════════
heading("3. COMPANY PROFILE", 1)
heading("Tech Elecon Pvt. Ltd.", 2)
para(
    "TECH ELECON plays pivotal role in IT sector of 'Elecon group' companies. Ever since its "
    "inception TECH ELECON has grown by leaps and bounds. In 1990 TECH ELECON achieved "
    "status of an independent corporate entity making its initial presence felt in the area of software "
    "development.\n\n"
    "TECH ELECON Pvt. Ltd., headquartered in Gujarat, has transitioned from an in-house IT service "
    "provider for the Elecon Group into a comprehensive IT solutions enterprise. It delivers wide-ranging "
    "software services, system integration, network infrastructure solutions, ERP deployments, and custom "
    "web application development to modern businesses. The company is characterized by a firm commitment to "
    "technological advancement and software engineering excellence.\n\n"
    "Address: Anand Sojitra Road, Vallabh Vidyanagar 388120, Gujarat, India\n"
    "Website: www.techelecon.com"
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 4: PROJECT OVERVIEW
# ═══════════════════════════════════════════
heading("4. PROJECT OVERVIEW", 1)
heading("4.1 TEAM INFORMATION", 2)
para("Developer: [Candidate Name]")
para("Organization: Tech Elecon Pvt. Ltd.")

heading("4.2 DIVISION OF WORK", 2)
para("The project work was systematically divided into frontend engineering (HTML/Bootstrap, Glassmorphism AI UI), "
     "backend engineering (Django, REST Framework, JWT), API integrations (Google Gemini, Mistral AI, Piston "
     "Sandbox), and Real-time Communications (WebRTC, Django Channels).")

heading("4.3 PROJECT DEFINITION", 2)
para("SmartRecruit is an end-to-end recruitment management system offering AI-based candidate screening, "
     "competency mapping, dynamically generated assessments, real-time proctored technical evaluations, and "
     "in-depth predictive reporting for HR practitioners.")

heading("4.4 PURPOSE & OBJECTIVE OF THE PROJECT", 2)
bullet("Purpose: Automate the labor-intensive stages of talent acquisition, reducing bias and costs.")
bullet("Objective 1: Reduce Shortlisting Time via Semantic Embeddings and Blockchain credential logging.")
bullet("Objective 2: Fully AI-driven Technical and HR evaluations without human interviewers.")

heading("4.5 SCOPE OF THE SYSTEM", 2)
para("The system strictly covers digital recruitment parameters: Recruiter job postings, Candidate application "
     "management, 4-tier assessment executions, Web Notifications (Toast), PWA install-ability, and Offer "
     "generation. Payroll and post-onboarding functions fall outside the scope.")

heading("4.6 KEY FEATURES", 2)
bullet("Multi-lingual Resume Match Analytics (Gemini/Tesseract)")
bullet("Secure 'Thunder Toast' UI Notifications for live status tracking")
bullet("Real-Time Webcam stream AI Emotion/Sentiment tracking")
bullet("PWA Native Support with offline architecture")
bullet("Integrated Code Editor backed by Piston Sandbox")

heading("4.7 SYSTEM DESIGN TECHNOLOGY STACK", 2)
bullet("Frontend: HTML5, CSS3 (Glassmorphism), Vanilla JavaScript, Bootstrap 5")
bullet("Backend: Python, Django 5.x, Django Channels")
bullet("Database: MySQL/PostgreSQL")
bullet("AI/ML: DeepFace, Gemini 1.5 API, Hugging Face Fallback")

heading("4.8 LIMITATIONS", 2)
bullet("Cannot replace the physical intuition required for culture-fit analysis of C-Level executives.")
bullet("Requires constant stable internet for WebRTC and AI Sandbox interaction.")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 5: SYSTEM ANALYSIS
# ═══════════════════════════════════════════
heading("5. SYSTEM ANALYSIS", 1)
heading("5.1 EXISTING SYSTEM", 2)
para("Current ATS (Applicant Tracking Systems) largely rely on basic Keyword Density matching. They lack semantic "
     "intent detection, provide no built-in proctoring for external tests, and entirely lack integrated AI Interviews. "
     "HRs manually review thousands of profiles resulting in extreme human error.")

heading("5.2 NEEDS OF SYSTEM", 2)
para("An intelligent solution that combines the tracking capabilities of an ATS with the evaluating capabilities "
     "of Senior Engineers and HR Specialists, bundled in a scalable and highly professional Glassmorphic UI.")

heading("5.3 PROPOSED SYSTEM", 2)
para("SmartRecruit delegates technical assessment and initial HR screening entirely to AI. It generates live reports, "
     "calculates 'Confidence' scores using real-time DeepFace tracking, and provides a unified dashboard for recruiters "
     "to instantly filter top performers.")

heading("5.4 MODULES", 2)
bullet("Account & Profile Module")
bullet("Recruiter Core (Post Job, Review Matches)")
bullet("Candidate Core (Apply, Status Tracking, Notifications)")
bullet("AI Assessment Engine (R1 & R2 gen)")
bullet("Live Interview Room (Piston IDE, WebRTC, DeepFace)")
bullet("Reporting & Offer Negotiations")

heading("5.5 PROJECT MANAGEMENT", 2)
heading("Feasibility Study", 3)
para("Technical, Operational, and Economic feasibility analyzed. Confirmed deployment via Docker on Linux environments is viable.")

heading("5.6 SYSTEM REQUIREMENT SPECIFICATION", 2)
bullet("Hardware: Dual-Core CPU, 4GB RAM minimum.")
bullet("Software: Python 3.11, Django, MySQL Server.")

heading("5.7 SOFTWARE DEVELOPMENT TOOLS", 2)
para("Visual Studio Code, GitHub for source tracking, MySQL Workbench, XAMPP for local testing.")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 6: SYSTEM DESIGN
# ═══════════════════════════════════════════
heading("6. SYSTEM DESIGN", 1)

heading("6.1 USE CASE DIAGRAM", 2)
add_image("use_case.png", 5.0, "Use Case Diagram")

heading("6.2 ACTIVITY DIAGRAMS", 2)
add_image("activity_flow.png", 5.0, "Activity Flow Diagram")

heading("6.3 ER DIAGRAM", 2)
add_image("er_diagram.png", 5.0, "Entity Relationship Diagram")

heading("6.4 DFD (DATA FLOW DIAGRAM)", 2)
add_image("dfd_level_0.png", 5.5, "Level 0 DFD - Context Diagram")
add_image("dfd_level_1.png", 5.5, "Level 1 DFD")
add_image("dfd_level_2.png", 5.5, "Level 2 DFD")

heading("6.5 SEQUENTIAL DIAGRAM", 2)
add_image("sequence_diagram.png", 5.0, "Sequence Diagram")

heading("6.6 CLASS DIAGRAM", 2)
add_image("class_diagram.png", 5.0, "Class Diagram")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 7: DATA DICTIONARY
# ═══════════════════════════════════════════
heading("7. DATA DICTIONARY", 1)

add_data_dictionary_table("user_users", [
    [1, "id", "INT", "Primary Key, Auto-Increment"],
    [2, "password", "VARCHAR(128)", "Not Null"],
    [3, "last_login", "DATETIME", "Null"],
    [4, "is_superuser", "TINYINT(1)", "Not Null, Default 0"],
    [5, "username", "VARCHAR(150)", "Unique, Not Null"],
    [6, "email", "VARCHAR(254)", "Not Null"],
    [7, "role", "VARCHAR(50)", "Not Null"]
])

add_data_dictionary_table("jobs_jobposting", [
    [1, "id", "INT", "Primary Key, Auto-Increment"],
    [2, "title", "VARCHAR(255)", "Not Null"],
    [3, "description", "LONGTEXT", "Not Null"],
    [4, "requirements", "LONGTEXT", "Not Null"],
    [5, "status", "VARCHAR(20)", "Not Null"],
    [6, "created_at", "DATETIME", "Not Null"],
    [7, "recruiter_id", "INT", "Foreign Key (users.id)"]
])

add_data_dictionary_table("jobs_application", [
    [1, "id", "INT", "Primary Key, Auto-Increment"],
    [2, "resume", "VARCHAR(100)", "Not Null"],
    [3, "match_score", "DOUBLE", "Null"],
    [4, "status", "VARCHAR(20)", "Not Null"],
    [5, "job_id", "INT", "Foreign Key (jobposting.id)"],
    [6, "candidate_id", "INT", "Foreign Key (users.id)"]
])

add_data_dictionary_table("jobs_interviewresult", [
    [1, "id", "INT", "Primary Key, Auto-Increment"],
    [2, "code_score", "DOUBLE", "Not Null"],
    [3, "sentiment_score", "DOUBLE", "Not Null"],
    [4, "feedback", "LONGTEXT", "Not Null"],
    [5, "application_id", "INT", "Foreign Key (application.id)"]
])
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 8: SYSTEM OUTPUT
# ═══════════════════════════════════════════
heading("8. SYSTEM OUTPUT", 1)

heading("8.1 SCREEN SHOT [ADMIN]", 2)
# Here we'll generate some placeholder or actual images depending on the screenshots available.
add_image("screenshot_landing.png", 5.5, "Landing Interface")

heading("8.2 SCREEN SHOT [USER]", 2)
add_image("screenshot_jobs.png", 5.5, "Available Job Listings")
add_image("screenshot_register.png", 5.5, "Candidate Registration")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 9: TESTING
# ═══════════════════════════════════════════
heading("9. TESTING", 1)
para("The Testing Plan outlines the strategy and steps followed to ensure the Real Estate/Recruitment Management System meets all requirements.")
bullet("Unit Testing: Authentication logic tested per module.")
bullet("Integration Testing: JWT middleware and AI API gateways evaluated.")
bullet("System Testing: Verified concurrent code submissions to Piston API.")
bullet("Security Testing: Checked WebSockets for hijacking risks.")

heading("Test Cases", 2)
add_table(
    ["Test Case", "Description", "Expected Status"],
    [
        ["TC01", "Invalid Login details", "Password error response received"],
        ["TC02", "Upload Malformed Resume", "System denies parse gracefully"],
        ["TC03", "Compile code in Sandbox", "Standard Output or Error string generated"],
    ],
    [1.0, 3.5, 1.5]
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 10: FUTURE ENHANCEMENT
# ═══════════════════════════════════════════
heading("10. FUTURE ENHANCEMENT", 1)
numbered("Integration with LinkedIn mapping and GitHub Code verification APIs.")
numbered("Deepfake analysis layer to detect AI-generated video imposters in real-time.")
numbered("Expanded AI Preparatory Labs combining VR and audio for interview prep.")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 11: CONCLUSION
# ═══════════════════════════════════════════
heading("11. CONCLUSION", 1)
para("The SmartRecruit Platform provides an incredibly cohesive connection between algorithmic evaluation "
     "and human resource management. By digitizing 90% of the manual effort in early-stage interviews, "
     "the system allows HR experts to focus strictly on human nuances, dramatically boosting overall productivity "
     "and ensuring scalable processing metrics without sacrificing software quality or candidate security.")

# ═══════════════════════════════════════════
#  CHAPTER 12: BIBLIOGRAPHY
# ═══════════════════════════════════════════
heading("12. BIBLIOGRAPHY", 1)
para("For the successful working of my project, I have referred to many sources for code snippets and logic:")
numbered("docs.djangoproject.com")
numbered("ai.google.dev (Gemini API)")
numbered("stackoverflow.com")
numbered("github.com/engineer-man/piston")

# ═══════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════
out_path = os.path.join(BASE_DIR, "SmartRecruit_Final_Documentation_V3.docx")
doc.save(out_path)
print(f"[OK] New DOCX saved successfully to {out_path}")
