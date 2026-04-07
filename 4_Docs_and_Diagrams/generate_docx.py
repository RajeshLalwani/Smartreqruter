"""
SmartRecruit AI Platform — Comprehensive DOCX Report Generator
Generates a professional, section-rich project documentation file.
Run: python generate_docx.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG = lambda name: os.path.join(BASE_DIR, name)

doc = Document()

# ── PAGE MARGINS ─────────────────────────────
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ── STYLE HELPERS ─────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_image(path, width=5.5, caption=None):
    full = IMG(path) if not os.path.isabs(path) else path
    if os.path.exists(full):
        doc.add_picture(full, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(9)
    else:
        p = doc.add_paragraph(f"[Figure: {caption or path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        # header background
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ═══════════════════════════════════════════
#  PAGE 1: TITLE PAGE
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph("SMARTRECRUIT AI RECRUITMENT PLATFORM")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.runs[0]
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 70, 127)

doc.add_paragraph()
sub = doc.add_paragraph("A Next-Generation AI-Powered Talent Acquisition Ecosystem")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
sub.runs[0].font.size = Pt(14)

doc.add_paragraph()
para("Project Report submitted as partial fulfilment of the degree of\nMaster of Computer Applications (MCA)", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_table(
    ["Detail", "Information"],
    [
        ["Project Title", "SmartRecruit AI Recruitment Platform"],
        ["Technology", "Python | Django | Google Gemini AI | PostgreSQL"],
        ["Department", "Post Graduate Dept. of Computer Science & Technology"],
        ["University", "Sardar Patel University, Vallabh Vidyanagar"],
        ["Academic Year", "2024–2026"],
        ["Company", "IR Info Tech Pvt. Ltd."],
    ],
    [2.0, 4.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════
#  CERTIFICATE
# ═══════════════════════════════════════════
heading("CERTIFICATE", 1)
para(
    "This is to certify that the project entitled \"SmartRecruit AI Recruitment Platform\" has been successfully "
    "carried out and completed by the students of Master of Computer Applications (MCA) as a partial "
    "fulfilment of their degree requirements under the Post Graduate Department of Computer Science and "
    "Technology, Sardar Patel University, Vallabh Vidyanagar, for the academic year 2024–2026.\n\n"
    "The project represents original work and has not been submitted for any other examination or degree."
)
doc.add_paragraph()
para("Guide Signature: __________________      Date: _______________")
para("HOD Signature: ___________________      Stamp: _______________")
doc.add_page_break()

# ═══════════════════════════════════════════
#  ABSTRACT
# ═══════════════════════════════════════════
heading("ABSTRACT", 1)
para(
    "SmartRecruit is a state-of-the-art, AI-driven recruitment platform engineered to modernize and "
    "completely automate the full talent acquisition lifecycle. Built on the Django framework with "
    "a multi-engine AI backbone (Google Gemini, Groq, Hugging Face), it eliminates the manual "
    "bottlenecks of traditional HR processes through intelligent automation.\n\n"
    "The system implements a 4-Round elimination pipeline: (R1) AI-generated Aptitude assessments, "
    "(R2) Domain-Specific Practical MCQs, (R3) Real-time AI Technical Interview with voice, facial "
    "sentiment analysis, and integrated coding challenges, and (R4) AI-facilitated Behavioral HR "
    "Interview via the 'Botanist' engine. The platform includes a live biometric proctoring engine, "
    "multi-lingual RAG-based resume parsing, blockchain credential verification, gamified coding "
    "arenas, salary negotiation AI, and a comprehensive recruiter analytics dashboard."
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  INDEX
# ═══════════════════════════════════════════
heading("TABLE OF CONTENTS", 1)
toc = [
    ("1", "Introduction & Project Background", ""),
    ("2", "Problem Statement & Objectives", ""),
    ("3", "System Features & Modules", ""),
    ("4", "Technology Stack", ""),
    ("5", "System Architecture & Design", ""),
    ("  5.1", "Use Case Diagram", ""),
    ("  5.2", "Data Flow Diagrams (Level 0, 1, 2)", ""),
    ("  5.3", "Activity Diagram", ""),
    ("  5.4", "Class Diagram", ""),
    ("  5.5", "Sequence Diagram", ""),
    ("6", "Database Design", ""),
    ("7", "AI & Machine Learning Modules", ""),
    ("8", "Security & Compliance", ""),
    ("9", "System Outputs (UI Screenshots)", ""),
    ("10", "Conclusion & Future Scope", ""),
]
for no, title, pg in toc:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{no}. ")
    r1.bold = True
    p.add_run(title)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════
heading("Chapter 1: Introduction & Project Background", 1)
para(
    "The modern recruitment landscape is overwhelmed by manual inefficiencies. HR professionals spend "
    "an estimated 70% of their time screening resumes — a task that is inherently biased and time-consuming. "
    "SmartRecruit addresses this challenge head-on by deploying Artificial Intelligence at every stage of "
    "the hiring funnel, from the first resume submission to the final offer letter generation."
)
heading("1.1 Project Background", 2)
para(
    "This project was developed as an industry-level internship deliverable at IR Info Tech Pvt. Ltd. "
    "The platform is designed around a real-world recruitment workflow, taking into consideration the needs "
    "of both growing startups and large enterprises. It was built using agile methodology over a 6-month "
    "period with regular stakeholder feedback loops."
)
heading("1.2 Scope", 2)
bullet("Complete digitization of the hire-to-onboard pipeline")
bullet("AI-powered resume parsing with multilingual support (30+ languages)")
bullet("Automated round-based candidate evaluation and elimination")
bullet("Real-time interview with AI interviewer, proctoring, and code execution")
bullet("Admin, Recruiter, and Candidate role-based access control")
bullet("Progressive Web App (PWA) support for mobile access")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 2: PROBLEM STATEMENT
# ═══════════════════════════════════════════
heading("Chapter 2: Problem Statement & Objectives", 1)
heading("2.1 Problem Statement", 2)
para(
    "Traditional recruitment systems are plagued by: manual resume screening bias, excessive time-to-hire "
    "(average 42 days globally), lack of objective candidate scoring, and inability to scale to high-volume "
    "application cycles. SmartRecruit is built to solve each of these pain points through AI automation."
)
heading("2.2 Objectives", 2)
numbered("Reduce time-to-shortlist by 85% using AI resume parsing")
numbered("Eliminate screening bias via blind-hiring toggles and objective AI scoring")
numbered("Automate technical skill evaluation through AI-generated assessments")
numbered("Enable real-time, scalable AI interviews with sentiment analysis")
numbered("Provide recruiters with actionable, predictive analytics dashboards")
numbered("Deliver a seamless, mobile-first Progressive Web App experience")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 3: FEATURES
# ═══════════════════════════════════════════
heading("Chapter 3: System Features & Modules", 1)

add_table(
    ["Module", "Key Features"],
    [
        ["Authentication & SSO", "JWT tokens, Azure AD, Okta SAML SSO, PWA install"],
        ["AI Resume Parser", "Multi-lingual RAG, TF-IDF + Semantic Embeddings, Blockchain verification"],
        ["Assessment Engine (R1/R2)", "AI-generated MCQs (Aptitude + Practical), Timed, Anti-cheat"],
        ["AI Interview Engine (R3)", "Voice STT/TTS, Coding Arena, Sentiment Analysis, Proctoring"],
        ["Botanist HR Engine (R4)", "Behavioral STAR-method evaluation, Multi-lingual voice"],
        ["Offer & Onboarding", "PDF Offer Generation, Salary Negotiation AI, 30-60-90 Day Roadmap"],
        ["Recruiter Analytics", "Hiring funnel charts, Retention prediction, Bias detection"],
        ["Admin Panel", "User management, Broadcast notifications, Platform settings"],
        ["AI Prep Lab", "50 Technical + 50 HR questions per role, confidence scoring"],
        ["Coding Arena", "Multi-language IDE, XP gamification, Leaderboards"],
        ["Notification System", "Branded HTML emails (async), in-app bell notifications, Webhook (n8n)"],
        ["PWA Support", "Offline cache, Install button, Service Worker, Push notifications"],
    ],
    [2.5, 4.0]
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 4: TECH STACK
# ═══════════════════════════════════════════
heading("Chapter 4: Technology Stack", 1)

add_table(
    ["Layer", "Technology", "Purpose"],
    [
        ["Backend Framework", "Python 3.11 / Django 4.x", "MVT architecture, ORM, security"],
        ["API Layer", "Django REST Framework + JWT", "JSON endpoints, token auth"],
        ["Real-time", "Django Channels + Redis", "WebSocket for live interview & proctoring"],
        ["Primary AI", "Google Gemini 1.5 Flash (google.genai)", "Interview, resume analysis, JD generation"],
        ["AI Fallback 1", "Groq API (LLaMA3-8B)", "Ultra-low latency voice & chat fallback"],
        ["AI Fallback 2", "Hugging Face (Mistral-7B)", "Ultimate offline fallback"],
        ["Embeddings", "text-embedding-004 / MiniLM-L12-v2", "Multilingual RAG vector matching"],
        ["Database", "SQLite (dev) / PostgreSQL (prod)", "Relational, ACID-compliant"],
        ["Frontend", "HTML5 + Vanilla JS + Bootstrap 5", "SSR templates, no heavy framework"],
        ["CSS System", "Custom enhancements.css (4000+ lines)", "Dark/Light theme CSS variables"],
        ["Authentication", "django-allauth + Azure AD + Okta SAML", "SSO enterprise login"],
        ["Static Files", "WhiteNoise + AWS S3 (optional)", "CDN-ready static serving"],
        ["Email", "Django SMTP + threading", "Async branded HTML emails"],
        ["Blockchain", "SHA-256 hashing ledger", "Resume/credential tamper-proof verification"],
        ["Deployment", "Gunicorn + Nginx / Cloud Run", "Production WSGI + reverse proxy"],
        ["PWA", "Service Worker + Web Manifest", "Offline-first mobile experience"],
    ],
    [1.7, 2.3, 2.5]
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 5: SYSTEM DESIGN
# ═══════════════════════════════════════════
heading("Chapter 5: System Architecture & Design", 1)
para(
    "The following section presents all major UML and DFD diagrams conforming to standard "
    "notation rules. All diagrams were generated from direct system analysis of the codebase."
)

heading("5.1 Use Case Diagram", 2)
para("Models the primary platform actors — Candidate, Recruiter, and Admin — "
     "with 33 documented use cases across the full hiring lifecycle.")
add_image("use_case.png", 5.5, "Figure 5.1 — Use Case Diagram (33 use cases, 3 actors)")
doc.add_page_break()

heading("5.2 Data Flow Diagrams", 2)
heading("Level 0 — Context Diagram", 3)
para("Shows the overall system boundary with all 6 external entities (Recruiter, Candidate, Admin, "
     "Gemini AI, n8n Webhooks, Blockchain Ledger).")
add_image("dfd_level_0.png", 5.5, "Figure 5.2.1 — DFD Level 0: Context Diagram")

heading("Level 1 — System Processes", 3)
para("Decomposes the SmartRecruit system into 7 primary sub-processes with their data stores.")
add_image("dfd_level_1.png", 5.5, "Figure 5.2.2 — DFD Level 1: Major Sub-Processes")

heading("Level 2 — Resume Intelligence Process", 3)
para("Detailed decomposition of the AI Resume Parsing module including multi-lingual detection, "
     "embedding, cosine similarity scoring, and threshold decision.")
add_image("dfd_level_2.png", 5.5, "Figure 5.2.3 — DFD Level 2: Resume Intelligence Detail")
doc.add_page_break()

heading("5.3 Activity Diagram", 2)
para("Captures the full 4-Round Hiring Pipeline including all decision gates, status transitions, "
     "and email notification triggers.")
add_image("activity_flow.png", 5.0, "Figure 5.3 — Activity Diagram: 4-Round Hiring Pipeline")
doc.add_page_break()

heading("5.4 Class Diagram", 2)
para("UML Class Diagram depicting 9 core persistent entities with their attributes, methods, "
     "and associations (including multiplicities).")
add_image("class_diagram.png", 5.5, "Figure 5.4 — Class Diagram: Core Data Models")
doc.add_page_break()

heading("5.5 Sequence Diagram", 2)
para("End-to-end interaction diagram tracing a candidate through the full hiring lifecycle — "
     "from application submission to offer acceptance — showing all async email triggers and AI calls.")
add_image("sequence_diagram.png", 5.5, "Figure 5.5 — Sequence Diagram: End-to-End Hiring Flow")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 6: DATABASE
# ═══════════════════════════════════════════
heading("Chapter 6: Database Design", 1)
para("The SmartRecruit database consists of 35+ inter-related models across 2 Django apps "
     "(core, jobs). Key entities and their relationships are documented below.")

add_table(
    ["Model", "App", "Key Fields", "Relationships"],
    [
        ["User", "core", "email, is_recruiter, is_staff, blind_hiring", "1→M JobPosting, 1→1 Candidate"],
        ["JobPosting", "jobs", "title, skills, salary_range, deadline, status", "1→M Application, M→M Technology"],
        ["Candidate", "jobs", "full_name, skills_extracted, resume_file, ai_score", "1→M Application"],
        ["Application", "jobs", "status (22 values), ai_score, technical_score", "M→1 Job, M→1 Candidate"],
        ["Assessment", "jobs", "test_type, score, passed, details (JSON)", "M→1 Application"],
        ["Interview", "jobs", "interview_type, ai_confidence_score, code_final", "M→1 Application, 1→M SentimentLog"],
        ["Offer", "jobs", "salary_offered, joining_date, status", "1→1 Application"],
        ["Notification", "jobs", "title, message, type, is_read", "M→1 User"],
        ["NegotiationSession", "jobs", "current_offer, candidate_counter, chat_history", "1→1 Offer"],
        ["CodingChallenge", "jobs", "title, difficulty, xp_reward, test_cases (JSON)", "1→M CodingSubmission"],
        ["OnboardingRoadmap", "jobs", "tasks (JSON), completion_percentage", "1→1 Application"],
        ["ProctoringLog", "jobs", "violation_type, severity, frame, timestamp", "M→1 Application"],
    ],
    [1.5, 0.8, 2.0, 2.2]
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 7: AI MODULES
# ═══════════════════════════════════════════
heading("Chapter 7: AI & Machine Learning Modules", 1)

heading("7.1 Three-Engine AI Architecture", 2)
para("SmartRecruit implements a graceful-degradation, 3-tier AI engine "
     "(AIEngine class in core/ai_engine.py):")
bullet("Primary: Google Gemini 1.5 Flash — high context, vision-capable (google.genai SDK)")
bullet("Fallback 1: Groq API (LLaMA3-8B-8192) — ultra-low latency for voice/chat")
bullet("Fallback 2: Hugging Face Inference API (Mistral-7B) — ultimate offline fallback")
para("If Gemini hits a rate limit, traffic is instantly rerouted to Groq, then HF — zero-downtime.")

heading("7.2 Resume Intelligence (RAG Pipeline)", 2)
add_table(
    ["Step", "Technology", "Output"],
    [
        ["1. Text Extraction", "pdfplumber / python-docx", "Raw resume text"],
        ["2. Language Detection", "langdetect", "ISO language code (hi, es, gu...)"],
        ["3. Semantic Embedding", "text-embedding-004 / MiniLM-L12-v2", "768-dim vector"],
        ["4. Cosine Similarity", "numpy / sklearn", "Match score 0–100"],
        ["5. Threshold Gate", "Business logic (≥60 shortlist)", "RESUME_SELECTED / REJECTED"],
    ],
    [1.2, 2.5, 2.8]
)

heading("7.3 AI Interview Engine", 2)
bullet("Voice STT via Web Speech API (browser-native, no API cost)")
bullet("Dynamic question generation by AIInterviewer based on job role + skill gaps")
bullet("Per-answer analysis: confidence score, technical accuracy, communication grade")
bullet("Live facial sentiment analysis via face-api.js and OpenCV proctoring")
bullet("Integrated Monaco code editor for live coding challenges within interview")

heading("7.4 Botanist AI Interviewer (HR Round)", 2)
bullet("Persona: 'The Botanist' — specialized in STAR-method behavioral evaluation")
bullet("Multi-lingual interview: candidate can speak in Hindi, Gujarati, Spanish, etc.")
bullet("Evaluation report always generated in English for recruiter dashboard")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 8: SECURITY
# ═══════════════════════════════════════════
heading("Chapter 8: Security & Compliance", 1)

add_table(
    ["Security Feature", "Implementation"],
    [
        ["CSRF Protection", "Django's built-in CSRF middleware on all POST forms"],
        ["XSS Prevention", "Django template auto-escaping, Content-Security-Policy headers"],
        ["SQL Injection", "ORM-based queries only, no raw SQL strings"],
        ["Authentication", "Bcrypt password hashing, JWT tokens (15-min expiry), SSO"],
        ["Rate Limiting", "Custom middleware + Nginx rate limiting"],
        ["Credential Verification", "SHA-256 blockchain hash for education/experience claims"],
        ["Interview Integrity", "Face detection proctoring, tab-switch detection, flag system"],
        ["HTTPS Enforcement", "SECURE_SSL_REDIRECT, HSTS 1-year max-age"],
        ["File Upload Safety", "MIME type + extension whitelist (PDF, DOC, DOCX only, 5MB cap)"],
        ["API Security", "DRF JWT authentication, CORS whitelist, rate throttling"],
    ],
    [2.5, 4.0]
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 9: UI SCREENSHOTS
# ═══════════════════════════════════════════
heading("Chapter 9: System Outputs (UI Screenshots)", 1)
para("The following screenshots demonstrate the deployed SmartRecruit platform. The system "
     "features a premium glassmorphism design with full Dark/Light theme switching capability.")

heading("9.1 Landing Page", 2)
add_image("screenshot_landing.png", 5.5, "Figure 9.1 — Landing Page (Dark Theme)")

heading("9.2 Neural Job Matrix", 2)
add_image("screenshot_jobs.png", 5.5, "Figure 9.2 — Job Listings with AI Match Scores")

heading("9.3 Candidate Registration", 2)
add_image("screenshot_register.png", 5.5, "Figure 9.3 — Candidate Registration Form")
doc.add_page_break()

# ═══════════════════════════════════════════
#  CHAPTER 10: CONCLUSION
# ═══════════════════════════════════════════
heading("Chapter 10: Conclusion & Future Scope", 1)
heading("10.1 Conclusion", 2)
para(
    "SmartRecruit successfully demonstrates that AI can be applied end-to-end across the complete "
    "talent acquisition lifecycle. The platform reduces time-to-shortlist from days to seconds, "
    "eliminates human bias from the screening process, and provides both candidates and recruiters "
    "with a world-class, data-driven experience. The system has been battle-tested through "
    "multiple iterations with zero system check errors and a fully functioning 4-round pipeline."
)

heading("10.2 Future Scope", 2)
numbered("Integration with LinkedIn, Naukri.com, and Indeed for automated job syndication")
numbered("Video proctoring with advanced anti-spoofing (3D face liveness detection)")
numbered("Reinforcement Learning-based question difficulty adaptation per candidate")
numbered("Multi-tenant SaaS architecture for enterprise white-labelling")
numbered("Advanced predictive HR analytics using historical hiring data regression models")
numbered("GDPR-compliant data anonymization and right-to-erasure automation")

doc.add_page_break()

# ═══════════════════════════════════════════
#  REFERENCES
# ═══════════════════════════════════════════
heading("References", 1)
numbered("Django Software Foundation. (2024). Django Documentation. https://docs.djangoproject.com/")
numbered("Google DeepMind. (2024). Gemini API Documentation. https://ai.google.dev/")
numbered("Hugging Face. (2024). Transformers Library. https://huggingface.co/docs/transformers/")
numbered("Bootstrap. (2024). Bootstrap 5 Documentation. https://getbootstrap.com/docs/5.3/")
numbered("OWASP. (2024). Top 10 Web Application Security Risks. https://owasp.org/")
numbered("Groq. (2024). Groq API Reference. https://console.groq.com/docs/")

# ═══════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════
out = os.path.join(os.path.dirname(__file__), "SmartRecruit_Final_Documentation_Compliant_v2.docx")
doc.save(out)
print(f"[OK] DOCX saved -> {out}")
