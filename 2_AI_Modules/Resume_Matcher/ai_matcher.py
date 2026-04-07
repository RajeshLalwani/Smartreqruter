import os
import re
import json
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
from google import genai
from google.genai import types

class ResumeMatcher:
    def __init__(self):
        self.vectorizer = TfidfVectorizer(stop_words='english')
        self.encoder = SentenceTransformer('all-MiniLM-L6-v2')
        self.api_key = os.environ.get("GEMINI_API_KEY")
        if self.api_key:
            self.client = genai.Client(api_key=self.api_key)
        else:
            self.client = None

    def extract_text(self, file_path):
        """Extracts text from PDF or DOCX file."""
        ext = os.path.splitext(file_path)[1].lower()
        text = ""
        
        try:
            if ext == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + " "
            elif ext == '.docx':
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + " "
                    
            elif ext == '.doc':
                try:
                    # Sometimes users rename .docx to .doc incorrectly
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        text += para.text + " "
                except Exception:
                    # Naive binary text extraction for older formats
                    with open(file_path, 'rb') as f:
                        b = f.read()
                        text = re.sub(r'[^\x20-\x7E]+', ' ', b.decode('latin-1', errors='replace'))

            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            
            return self.clean_text(text)
        except Exception as e:
            print(f"Error extracting text: {e}")
            return ""

    def clean_text(self, text):
        """Cleans text by removing special characters and extra spaces."""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s]', '', text)
        return text.lower().strip()

    def calculate_match(self, resume_text, job_description):
        """
        Calculates the cosine similarity between resume and job description.
        Returns a score out of 100.
        """
        if not resume_text or not job_description:
            return 0.0

        resume_vec = self.encoder.encode([resume_text])
        job_vec = self.encoder.encode([job_description])
        similarity_matrix = cosine_similarity(resume_vec, job_vec)
        
        score = similarity_matrix[0][0] * 100
        return round(score, 2)

    def get_semantic_score(self, resume_text, job_description):
        """
        [ELITE] Uses Gemini 2.0 Flash to evaluate deep semantic fit.
        """
        if not self.client:
            return None
            
        prompt = f"""
        You are an ELITE ATS (Applicant Tracking System) Auditor.
        Compare the candidate's Resume against the Job Description (JD).
        
        RESUME:
        {resume_text[:6000]}
        
        JOB DESCRIPTION:
        {job_description[:4000]}

        TASK:
        Provide a deep semantic analysis in JSON format:
        {{
            "semantic_score": <0-100 based on core tech alignment and seniority>,
            "relevance_summary": "<1-2 sentences on why they fit or don't fit>",
            "missing_critical_skills": ["<list of high-priority JD skills they lack>"],
            "bonus_skills": ["<valuable skills they have not mentioned in JD but useful for role>"]
        }}
        """
        try:
            response = self.client.models.generate_content(
                model='gemini-2.0-flash',
                config=types.GenerateContentConfig(response_mime_type='application/json'),
                contents=prompt
            )
            return json.loads(response.text)
        except Exception as e:
            print(f"Semantic Match Error: {e}")
            return None

    def extract_keywords(self, resume_text, job_description, top_n=10):
        """
        Extracts matched and missing keywords using TF-IDF.
        Returns a dict with 'matched' and 'missing' lists.
        """
        if not resume_text or not job_description:
            return {'matched': [], 'missing': []}

        documents = [resume_text, job_description]
        tfidf_matrix = self.vectorizer.fit_transform(documents)
        feature_names = self.vectorizer.get_feature_names_out() # or get_feature_names() for older sklearn
        
        # Get scores
        resume_scores = tfidf_matrix[0].toarray()[0]
        job_scores = tfidf_matrix[1].toarray()[0]
        
        # Create dicts {word: score}
        resume_dict = {feature_names[i]: resume_scores[i] for i in range(len(feature_names)) if resume_scores[i] > 0}
        job_dict = {feature_names[i]: job_scores[i] for i in range(len(feature_names)) if job_scores[i] > 0}
        
        # Find Matches and Missing
        matched = []
        missing = []
        
        # Sort job keywords by importance (score)
        sorted_job_keywords = sorted(job_dict.items(), key=lambda x: x[1], reverse=True)
        
        for word, score in sorted_job_keywords:
            if word in resume_dict:
                matched.append(word)
            else:
                missing.append(word)
                
        return {
            'matched': matched[:top_n],
            'missing': missing[:top_n]
        }
